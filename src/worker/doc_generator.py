"""Document Generator — PDF (ReportLab) and DOCX (python-docx).

PDF structure:
  Section 1: Per-city lists (hotels, attractions, activities, restaurants, bars)
  Section 2: Daily schedule table (Date | Accommodation | City | Morning | Afternoon | Evening)

DOCX structure:
  Section 1: Per-city lists with hyperlinks
  Section 2: Table with same columns as PDF

Both documents are written to BytesIO buffers — callers handle file storage.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any

import structlog
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
)

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.worker.trip_composer import ComposedTrip, DaySchedule
from src.worker.rules_engine import RecommendationCandidate

log = structlog.get_logger(__name__)

MAX_ITEMS_PER_TYPE = 10

# ── Shared helpers ────────────────────────────────────────────────────────────

RECOMMENDATION_TYPES_ORDER = ["hotel", "attraction", "activity", "restaurant", "bar"]
TYPE_LABELS = {
    "hotel": "Hotels",
    "attraction": "Attractions",
    "activity": "Activities",
    "restaurant": "Restaurants",
    "bar": "Bars & Nightlife",
}


# ── PDF Generation ────────────────────────────────────────────────────────────

def _build_pdf_city_section(
    city: str,
    city_recs: dict[str, list[RecommendationCandidate]],
    styles: Any,
) -> list[Any]:
    """Build ReportLab flowables for a single city's recommendation section."""
    elements: list[Any] = []

    elements.append(Paragraph(city, styles["Heading1"]))
    elements.append(Spacer(1, 0.2 * cm))

    for rec_type in RECOMMENDATION_TYPES_ORDER:
        items = city_recs.get(rec_type, [])[:MAX_ITEMS_PER_TYPE]
        if not items:
            continue

        label = TYPE_LABELS.get(rec_type, rec_type.title())
        elements.append(Paragraph(label, styles["Heading2"]))

        for rec in items:
            rating_str = f"⭐ {rec.rating}" if rec.rating else ""
            reviews_str = f"({rec.review_count} reviews)" if rec.review_count else ""
            price_str = f" | {rec.price_hint}" if rec.price_hint else ""

            line = (
                f"<b>{rec.name}</b> {rating_str} {reviews_str}{price_str} "
                f'— <a href="{rec.source_url}" color="blue">{rec.source_name}</a>'
            )
            elements.append(Paragraph(line, styles["Normal"]))
            elements.append(Spacer(1, 0.1 * cm))

        elements.append(Spacer(1, 0.3 * cm))

    elements.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey))
    elements.append(Spacer(1, 0.5 * cm))
    return elements


def _build_pdf_schedule_table(
    daily_schedule: list[DaySchedule],
    styles: Any,
) -> list[Any]:
    """Build ReportLab Table flowable for the daily schedule."""
    elements: list[Any] = []

    elements.append(Paragraph("Daily Itinerary", styles["Heading1"]))
    elements.append(Spacer(1, 0.3 * cm))

    header = ["Date", "City", "Accommodation", "Morning", "Afternoon", "Evening"]
    table_data: list[list[Any]] = [header]

    for day in daily_schedule:
        row = [
            Paragraph(day.date_str, styles["Normal"]),
            Paragraph(day.city, styles["Normal"]),
            Paragraph(day.accommodation, styles["Normal"]),
            Paragraph(day.morning, styles["Normal"]),
            Paragraph(day.afternoon, styles["Normal"]),
            Paragraph(day.evening, styles["Normal"]),
        ]
        table_data.append(row)

    col_widths = [2.5 * cm, 3.5 * cm, 4 * cm, 4.5 * cm, 4.5 * cm, 4.5 * cm]

    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F3F4")]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BDC3C7")),
                ("FONTSIZE", (0, 1), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    elements.append(table)
    return elements


def generate_pdf(trip: ComposedTrip, trip_meta: dict[str, Any]) -> bytes:
    """Generate the full PDF itinerary. Returns raw bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=f"Luxury Europe Trip — {trip_meta.get('country', '')}",
    )

    styles = getSampleStyleSheet()
    # Add a link-safe Normal style
    link_style = ParagraphStyle(
        "LinkNormal",
        parent=styles["Normal"],
        fontSize=9,
        leading=12,
    )
    styles.add(link_style)

    elements: list[Any] = []

    # ── Cover / Header ────────────────────────────────────────────────────────
    elements.append(Paragraph(
        f"Luxury Europe Trip — {trip_meta.get('country', '')}",
        styles["Title"],
    ))
    elements.append(Paragraph(
        f"{trip_meta.get('days', '')} days | "
        f"{trip_meta.get('party_size', '')} | "
        f"{trip_meta.get('dates_or_month', '')}",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 1 * cm))

    # ── Section 1: Per-city recommendations ──────────────────────────────────
    elements.append(Paragraph("Recommendations by City", styles["Heading1"]))
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#2C3E50")))
    elements.append(Spacer(1, 0.5 * cm))

    for slot in trip.city_slots:
        city_recs = trip.recommendations_by_city.get(slot.city, {})
        elements.extend(_build_pdf_city_section(slot.city, city_recs, styles))

    # ── Section 2: Daily Schedule ─────────────────────────────────────────────
    elements.append(Spacer(1, 1 * cm))
    elements.extend(_build_pdf_schedule_table(trip.daily_schedule, styles))

    doc.build(elements)
    buffer.seek(0)
    pdf_bytes = buffer.getvalue()

    log.info("pdf_generated", size_bytes=len(pdf_bytes))
    return pdf_bytes


# ── DOCX Generation ───────────────────────────────────────────────────────────

def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Add a hyperlink to a DOCX paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_docx_city_section(
    doc: DocxDocument,
    city: str,
    city_recs: dict[str, list[RecommendationCandidate]],
) -> None:
    """Add a city's recommendation section to the DOCX document."""
    doc.add_heading(city, level=1)

    for rec_type in RECOMMENDATION_TYPES_ORDER:
        items = city_recs.get(rec_type, [])[:MAX_ITEMS_PER_TYPE]
        if not items:
            continue

        label = TYPE_LABELS.get(rec_type, rec_type.title())
        doc.add_heading(label, level=2)

        for rec in items:
            para = doc.add_paragraph()
            run = para.add_run(rec.name)
            run.bold = True

            rating_str = f"  ⭐ {rec.rating}" if rec.rating else ""
            reviews_str = f" ({rec.review_count} reviews)" if rec.review_count else ""
            price_str = f" | {rec.price_hint}" if rec.price_hint else ""

            para.add_run(f"{rating_str}{reviews_str}{price_str}  ")
            _add_hyperlink(para, rec.source_url, rec.source_name)


def _add_docx_schedule_table(
    doc: DocxDocument,
    daily_schedule: list[DaySchedule],
) -> None:
    """Add the daily itinerary table to the DOCX document."""
    doc.add_heading("Daily Itinerary", level=1)

    headers = ["Date", "City", "Accommodation", "Morning", "Afternoon", "Evening"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for day in daily_schedule:
        row_cells = table.add_row().cells
        row_cells[0].text = day.date_str
        row_cells[1].text = day.city
        row_cells[2].text = day.accommodation
        row_cells[3].text = day.morning
        row_cells[4].text = day.afternoon
        row_cells[5].text = day.evening


def generate_docx(trip: ComposedTrip, trip_meta: dict[str, Any]) -> bytes:
    """Generate the full DOCX itinerary. Returns raw bytes."""
    doc = DocxDocument()

    # Title
    title = doc.add_heading(
        f"Luxury Europe Trip — {trip_meta.get('country', '')}",
        level=0,
    )

    doc.add_paragraph(
        f"{trip_meta.get('days', '')} days | "
        f"{trip_meta.get('party_size', '')} | "
        f"{trip_meta.get('dates_or_month', '')}"
    )

    # Section 1
    doc.add_heading("Recommendations by City", level=1)
    for slot in trip.city_slots:
        city_recs = trip.recommendations_by_city.get(slot.city, {})
        _add_docx_city_section(doc, slot.city, city_recs)
        doc.add_page_break()

    # Section 2
    _add_docx_schedule_table(doc, trip.daily_schedule)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.getvalue()

    log.info("docx_generated", size_bytes=len(docx_bytes))
    return docx_bytes
